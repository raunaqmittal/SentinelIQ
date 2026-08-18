"""PDF loading and text extraction."""

import hashlib
import logging
import uuid
from datetime import UTC, datetime
from functools import lru_cache
from pathlib import Path

import docx
import pymupdf

from sentineliq.components.models.schemas import LoadedDocument, PageSpan
from sentineliq.config import load_app_config
from sentineliq.exceptions import DocumentLoadError

logger = logging.getLogger(__name__)

PAGE_SEPARATOR = "\n"
MIN_TEXT_CHARS = 20

#: First bytes each supported format must start with. Checking the content is
#: the point — an extension is chosen by whoever supplies the file, so it is
#: not evidence of anything (NFR-003, Context.md 26.H). `.txt` has no signature
#: and is validated by decoding instead.
SIGNATURES = {".pdf": b"%PDF-", ".docx": b"PK\x03\x04"}


@lru_cache(maxsize=1)
def _configured_max_bytes() -> int:
    """The configured size limit, read once."""
    return load_app_config().ingestion.max_file_bytes


def validate_file(path: Path, max_bytes: int | None = None) -> None:
    """Reject a file before any parser touches it.

    Checks that the file exists, is within the configured size limit, has a
    supported extension, and that its **content** matches that extension.

    Raises:
        DocumentLoadError: The file fails any of those checks.
    """
    if not path.is_file():
        raise DocumentLoadError(f"No such document: {path}")

    limit = _configured_max_bytes() if max_bytes is None else max_bytes
    size = path.stat().st_size
    if size > limit:
        raise DocumentLoadError(f"{path.name} is {size} bytes, over the {limit} limit")

    suffix = path.suffix.lower()
    if suffix not in SIGNATURES and suffix != ".txt":
        raise DocumentLoadError(f"Unsupported file type: {path.name}")

    if suffix == ".txt":
        # No signature to check, so require it to be readable text. A NUL byte
        # means it is really a binary file wearing a .txt extension.
        try:
            body = path.read_bytes()
            body.decode("utf-8")
        except UnicodeDecodeError as error:
            raise DocumentLoadError(f"{path.name} is not valid UTF-8 text") from error
        if b"\x00" in body:
            raise DocumentLoadError(f"{path.name} is binary, not text")
        return

    with open(path, "rb") as handle:
        start = handle.read(len(SIGNATURES[suffix]))
    if start != SIGNATURES[suffix]:
        raise DocumentLoadError(f"{path.name} content is not a real {suffix} file")


def compute_sha256(path: Path) -> str:
    """Return the SHA-256 hex digest of a file."""
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _page_spans(page_texts: list[str]) -> list[PageSpan]:
    """Map each page onto its character range in the joined text."""
    spans = []
    start = 0
    for number, text in enumerate(page_texts, start=1):
        spans.append(PageSpan(page_number=number, start=start, end=start + len(text)))
        start += len(text) + len(PAGE_SEPARATOR)
    return spans


def load_pdf(path: Path, *, document_id: str | None = None) -> LoadedDocument:
    """Load a PDF into a LoadedDocument with page offsets into its text.

    Raises:
        DocumentLoadError: The file is missing, unparseable, or has no text.
    """
    validate_file(path)

    try:
        with pymupdf.open(path) as pdf:
            page_texts = [page.get_text("text") for page in pdf]
    except Exception as error:  # pymupdf raises assorted low-level errors
        raise DocumentLoadError(f"Failed to parse {path.name}") from error

    text = PAGE_SEPARATOR.join(page_texts)
    if len(text.strip()) < MIN_TEXT_CHARS:
        raise DocumentLoadError(f"{path.name} has no extractable text; may be scanned")

    document = LoadedDocument(
        document_id=document_id or str(uuid.uuid4()),
        document_name=path.name,
        text=text,
        pages=_page_spans(page_texts),
        sha256=compute_sha256(path),
        loaded_at=datetime.now(UTC),
    )
    logger.info(
        "Loaded document",
        extra={"document_id": document.document_id, "pages": len(page_texts)},
    )
    return document


def load_txt(path: Path, *, document_id: str | None = None) -> LoadedDocument:
    """Load a plain text file into a LoadedDocument.

    TXT files have no pages, so `pages` is always empty.

    Raises:
        DocumentLoadError: The file is missing or has no usable text.
    """
    validate_file(path)

    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as error:
        raise DocumentLoadError(f"{path.name} is not valid UTF-8 text") from error

    if len(text.strip()) < MIN_TEXT_CHARS:
        raise DocumentLoadError(f"{path.name} has no extractable text")

    document = LoadedDocument(
        document_id=document_id or str(uuid.uuid4()),
        document_name=path.name,
        text=text,
        pages=[],
        sha256=compute_sha256(path),
        loaded_at=datetime.now(UTC),
    )
    logger.info("Loaded document", extra={"document_id": document.document_id})
    return document


def load_docx(path: Path, *, document_id: str | None = None) -> LoadedDocument:
    """Load a Word document into a LoadedDocument.

    DOCX files have no pages, so `pages` is always empty.

    Raises:
        DocumentLoadError: The file is missing, unparseable, or has no text.
    """
    validate_file(path)

    try:
        paragraphs = [p.text for p in docx.Document(str(path)).paragraphs]
    except Exception as error:  # python-docx raises assorted low-level errors
        raise DocumentLoadError(f"Failed to parse {path.name}") from error

    text = "\n".join(paragraphs)
    if len(text.strip()) < MIN_TEXT_CHARS:
        raise DocumentLoadError(f"{path.name} has no extractable text")

    document = LoadedDocument(
        document_id=document_id or str(uuid.uuid4()),
        document_name=path.name,
        text=text,
        pages=[],
        sha256=compute_sha256(path),
        loaded_at=datetime.now(UTC),
    )
    logger.info("Loaded document", extra={"document_id": document.document_id})
    return document
